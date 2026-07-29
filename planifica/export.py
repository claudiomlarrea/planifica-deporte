from __future__ import annotations

import html
import json
from typing import Any

from planifica.surveys import list_configured_surveys
from planifica.utils import ensure_foda


def _line(label: str, value: str) -> str:
    v = (value or "").strip()
    return f"### {label}\n{v if v else '—'}\n\n"


def _esc(value: Any) -> str:
    return html.escape(str(value or "").strip() or "—")


def _block(label: str, value: Any) -> str:
    text = str(value or "").strip() or "—"
    if text == "—":
        paragraphs = "<p class='pd-empty'>—</p>"
    else:
        paragraphs = "".join(
            f"<p>{html.escape(line)}</p>" if line.strip() else "<br/>"
            for line in text.splitlines()
        )
    return f"""
    <div class="pd-field">
      <div class="pd-field-label">{html.escape(label)}</div>
      <div class="pd-field-body">{paragraphs}</div>
    </div>
    """


def plan_to_html(title: str, payload: dict[str, Any]) -> str:
    """Vista previa tipográfica del plan (HTML), no Markdown crudo."""
    org = payload.get("org") or {}
    pei = payload.get("pei") or {}
    pg = payload.get("proyecto_guia") or {}
    dafo = ensure_foda(payload)
    cim = payload.get("cimientos") or {}
    rend = payload.get("rendimiento") or {}
    rrhh = payload.get("rrhh") or {}
    vol = payload.get("voluntarios") or {}
    acciones = payload.get("acciones") or []

    pais = org.get("pais") or org.get("region") or "—"
    provincia = org.get("provincia") or ""
    ubicacion = f"{pais}" + (f" · {provincia}" if provincia else "")
    pei_titulo = pei.get("nombre") or title

    meta = f"""
    <div class="pd-meta">
      <div><span>Organización</span><strong>{_esc(org.get('nombre'))}</strong></div>
      <div><span>Tipo</span><strong>{_esc(org.get('tipo'))}</strong></div>
      <div><span>Ubicación</span><strong>{_esc(ubicacion)}</strong></div>
      <div><span>Horizonte</span><strong>{_esc(org.get('horizonte_anios'))} años</strong></div>
    </div>
    """

    pei_block = f"""
      <h3>Plan Estratégico Institucional</h3>
      {_block("Nombre del PEI", pei_titulo)}
      {_block("Período", pei.get("periodo"))}
      {_block("Versión", pei.get("version"))}
      {_block("Aprobado por", pei.get("aprobado_por"))}
      {_block("Fecha de aprobación", pei.get("fecha_aprobacion"))}
    """

    if acciones:
        cards = []
        n = 0
        for a in acciones:
            if not any(str(a.get(k) or "").strip() for k in ("accion", "responsable", "plazo", "kpi")):
                continue
            n += 1
            cards.append(
                f"""
                <div class="pd-action">
                  <div class="pd-action-title">{n}. {_esc(a.get('accion'))}</div>
                  <div class="pd-action-meta">
                    <span><b>Prioridad:</b> {_esc(a.get('prioridad'))}</span>
                    <span><b>Responsable:</b> {_esc(a.get('responsable'))}</span>
                    <span><b>Plazo:</b> {_esc(a.get('plazo'))}</span>
                    <span><b>KPI:</b> {_esc(a.get('kpi'))}</span>
                    <span><b>Recursos:</b> {_esc(a.get('recursos'))}</span>
                    <span><b>Estado:</b> {_esc(a.get('estado'))}</span>
                  </div>
                </div>
                """
            )
        acciones_html = "".join(cards) if cards else "<p class='pd-empty'>Sin acciones registradas.</p>"
    else:
        acciones_html = "<p class='pd-empty'>Sin acciones registradas.</p>"

    return f"""
    <div class="pd-doc">
      <div class="pd-doc-header">
        <div class="pd-doc-badge">Primer Plan Estratégico Institucional</div>
        <h2>{_esc(pei_titulo)}</h2>
      </div>
      {meta}
      {pei_block}
      <h3>1 · Análisis FODA</h3>
      <div class="pd-grid-2">
        {_block("Fortalezas", dafo.get("fortalezas"))}
        {_block("Debilidades", dafo.get("debilidades"))}
        {_block("Oportunidades", dafo.get("oportunidades"))}
        {_block("Amenazas", dafo.get("amenazas"))}
      </div>
      <h3>2 · Visión, misión y valores</h3>
      {_block("Visión", cim.get("vision"))}
      {_block("Misión", cim.get("mision"))}
      {_block("Valores", cim.get("valores"))}
      <h3>3 · Prioridades y objetivos</h3>
      {_block("Prioridades estratégicas", payload.get("prioridades"))}
      {_block("Objetivos SMART", payload.get("objetivos_smart"))}
      <h3>4 · Plan de acción</h3>
      {acciones_html}
      <h3>5 · Indicadores y evaluación</h3>
      {_block("Indicadores del plan de acción (KPI)", rend.get("kpis"))}
      {_block("Frecuencia de evaluación", rend.get("frecuencia_evaluacion"))}
      {_block("Informes al comité", rend.get("informes_comite"))}
      <h3>6 · Recursos humanos y voluntarios</h3>
      {_block("Roles clave", rrhh.get("roles_clave"))}
      {_block("Brechas de formación", rrhh.get("brechas_formacion"))}
      {_block("Reclutamiento", rrhh.get("reclutamiento"))}
      {_block("Necesidades de voluntariado", vol.get("necesidades"))}
      {_block("Motivaciones", vol.get("motivaciones"))}
      {_block("Formación", vol.get("formacion"))}
      {_block("Reconocimiento", vol.get("reconocimiento"))}
      <h3>7 · Primer proyecto derivado del PEI (opcional)</h3>
      {_block("Proyecto", pg.get("nombre"))}
      {_block("Objetivo", pg.get("objetivo"))}
      {_block("Contribuye al PEI en", pg.get("vinculo_estrategico"))}
      {_block("Criterios de éxito", pg.get("criterios_exito"))}
      <div class="pd-doc-footer">
        Primer PEI generado con Rumbo Deporte · Manual COI (2020), Unidades 53–57
      </div>
    </div>
    """


def plan_to_markdown(title: str, payload: dict[str, Any]) -> str:
    org = payload.get("org") or {}
    pei = payload.get("pei") or {}
    pg = payload.get("proyecto_guia") or {}
    dafo = ensure_foda(payload)
    cim = payload.get("cimientos") or {}
    rend = payload.get("rendimiento") or {}
    rrhh = payload.get("rrhh") or {}
    vol = payload.get("voluntarios") or {}
    pei_titulo = pei.get("nombre") or title

    lines = [
        f"# Primer Plan Estratégico Institucional — {pei_titulo}",
        "",
        f"**Organización:** {org.get('nombre', '')}",
        f"**Tipo:** {org.get('tipo', '')}",
        f"**País:** {org.get('pais') or org.get('region', '')}"
        + (f" · **Provincia:** {org.get('provincia')}" if org.get("provincia") else ""),
        f"**Horizonte:** {org.get('horizonte_anios', '')} años",
        "",
        "## Identidad del PEI",
        _line("Nombre del PEI", pei_titulo),
        _line("Período", pei.get("periodo", "")),
        _line("Versión", pei.get("version", "")),
        _line("Quién lo aprobará", pei.get("aprobado_por", "")),
        _line("Fecha de aprobación", pei.get("fecha_aprobacion", "")),
        "",
        "---",
        "",
        "## 1. Análisis FODA",
        _line("Fortalezas", dafo.get("fortalezas", "")),
        _line("Debilidades", dafo.get("debilidades", "")),
        _line("Oportunidades", dafo.get("oportunidades", "")),
        _line("Amenazas", dafo.get("amenazas", "")),
        "",
        "## 2. Visión, misión y valores",
        _line("Visión", cim.get("vision", "")),
        _line("Misión", cim.get("mision", "")),
        _line("Valores", cim.get("valores", "")),
        "",
        "## 3. Prioridades y objetivos",
        _line("Prioridades estratégicas", payload.get("prioridades", "")),
        _line("Objetivos SMART", payload.get("objetivos_smart", "")),
        "",
        "## 4. Plan de acción",
    ]

    acciones = payload.get("acciones") or []
    if acciones:
        for i, a in enumerate(acciones, 1):
            lines.append(
                f"{i}. **{a.get('accion', '')}** — Responsable: {a.get('responsable', '')} · "
                f"Plazo: {a.get('plazo', '')} · KPI: {a.get('kpi', '')} · "
                f"Recursos: {a.get('recursos', '')} · Estado: {a.get('estado', '')}"
            )
        lines.append("")
    else:
        lines.append("_Sin acciones registradas._\n")

    lines.extend(
        [
            "## 5. Evaluación e informes",
            _line("Indicadores del plan de acción (KPI)", rend.get("kpis", "")),
            _line("Frecuencia de evaluación", rend.get("frecuencia_evaluacion", "")),
            _line("Informes al comité", rend.get("informes_comite", "")),
            "",
            "## 6. Recursos humanos y voluntarios",
            _line("Roles clave", rrhh.get("roles_clave", "")),
            _line("Brechas de formación", rrhh.get("brechas_formacion", "")),
            _line("Reclutamiento", rrhh.get("reclutamiento", "")),
            _line("Necesidades de voluntariado", vol.get("necesidades", "")),
            _line("Motivaciones", vol.get("motivaciones", "")),
            _line("Formación", vol.get("formacion", "")),
            _line("Reconocimiento", vol.get("reconocimiento", "")),
            "",
            "## 7. Primer proyecto derivado del PEI (opcional)",
            _line("Proyecto", pg.get("nombre", "")),
            _line("Objetivo", pg.get("objetivo", "")),
            _line("Contribuye al PEI en", pg.get("vinculo_estrategico", "")),
            _line("Criterios de éxito", pg.get("criterios_exito", "")),
            "",
            "## 8. Actividades de ejecución",
        ]
    )
    actividades = payload.get("actividades") or []
    n = 0
    for a in actividades:
        if not str(a.get("titulo") or "").strip():
            continue
        n += 1
        lines.append(
            f"{n}. **{a.get('titulo', '')}** — Prioridad: {a.get('prioridad', '')} · "
            f"Estado: {a.get('estado', '')} · KPI: {a.get('kpi_nombre', '')} · "
            f"Meta: {a.get('meta', '')} · Avance: {a.get('avance', '')}"
        )
    if n == 0:
        lines.append("_Sin actividades registradas._")

    surveys = list_configured_surveys(payload)
    if surveys:
        lines.extend(["", "## 10. Encuestas de consulta"])
        for s in surveys:
            dest = f" · Destinatarios: {s['destinatarios']}" if s.get("destinatarios") else ""
            lines.append(f"- **{s['etiqueta']}**: {s['url']}{dest}")

    lines.extend(
        [
            "",
            "---",
            "",
            "_Primer PEI generado con Rumbo Deporte · Manual COI (2020), Unidades 53–57._",
        ]
    )
    return "\n".join(lines)


def plan_backup_bytes(title: str, payload: dict[str, Any]) -> bytes:
    doc = {"format": "rumbo-backup", "version": 1, "title": title, "payload": payload}
    return json.dumps(doc, ensure_ascii=False, indent=2).encode("utf-8")


def parse_plan_backup(raw: bytes) -> tuple[str, dict[str, Any]]:
    data = json.loads(raw.decode("utf-8"))
    fmt = data.get("format")
    if fmt not in {"rumbo-backup", "planifica-deporte-backup"}:
        raise ValueError("Archivo no reconocido como respaldo Rumbo Deporte.")
    return str(data.get("title") or "Plan importado"), data.get("payload") or {}


def _docx_para(doc, label: str, value: Any) -> None:
    from docx.shared import Pt

    text = str(value or "").strip() or "—"
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(11)
    body = p.add_run(text)
    body.font.size = Pt(11)


def plan_to_docx(title: str, payload: dict[str, Any]) -> bytes:
    """Genera un .docx del PEI para descargar en Word."""
    from io import BytesIO

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    org = payload.get("org") or {}
    pei = payload.get("pei") or {}
    pg = payload.get("proyecto_guia") or {}
    dafo = ensure_foda(payload)
    cim = payload.get("cimientos") or {}
    rend = payload.get("rendimiento") or {}
    rrhh = payload.get("rrhh") or {}
    vol = payload.get("voluntarios") or {}
    pei_titulo = pei.get("nombre") or title
    pais = org.get("pais") or org.get("region") or "—"
    provincia = org.get("provincia") or ""
    ubicacion = f"{pais}" + (f" · {provincia}" if provincia else "")

    doc = Document()
    h = doc.add_heading(pei_titulo, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x04, 0x4A, 0x30)

    sub = doc.add_paragraph("Primer Plan Estratégico Institucional · Rumbo Deporte")
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(10)

    doc.add_heading("Organización", level=1)
    _docx_para(doc, "Nombre", org.get("nombre"))
    _docx_para(doc, "Tipo", org.get("tipo"))
    _docx_para(doc, "Ubicación", ubicacion)
    _docx_para(doc, "Horizonte", f"{org.get('horizonte_anios', '')} años")

    doc.add_heading("Identidad del PEI", level=1)
    _docx_para(doc, "Nombre", pei_titulo)
    _docx_para(doc, "Período", pei.get("periodo"))
    _docx_para(doc, "Versión", pei.get("version"))
    _docx_para(doc, "Quién lo aprobará", pei.get("aprobado_por"))
    _docx_para(doc, "Fecha de aprobación", pei.get("fecha_aprobacion"))

    doc.add_heading("1. Análisis FODA", level=1)
    _docx_para(doc, "Fortalezas", dafo.get("fortalezas"))
    _docx_para(doc, "Debilidades", dafo.get("debilidades"))
    _docx_para(doc, "Oportunidades", dafo.get("oportunidades"))
    _docx_para(doc, "Amenazas", dafo.get("amenazas"))

    doc.add_heading("2. Visión, misión y valores", level=1)
    _docx_para(doc, "Visión", cim.get("vision"))
    _docx_para(doc, "Misión", cim.get("mision"))
    _docx_para(doc, "Valores", cim.get("valores"))

    doc.add_heading("3. Prioridades y objetivos", level=1)
    _docx_para(doc, "Prioridades", payload.get("prioridades"))
    _docx_para(doc, "Objetivos SMART", payload.get("objetivos_smart"))

    doc.add_heading("4. Plan de acción", level=1)
    acciones = payload.get("acciones") or []
    n = 0
    for a in acciones:
        if not any(str(a.get(k) or "").strip() for k in ("accion", "responsable", "plazo", "kpi")):
            continue
        n += 1
        doc.add_heading(f"{n}. {a.get('accion') or 'Acción'}", level=2)
        _docx_para(doc, "Prioridad", a.get("prioridad"))
        _docx_para(doc, "Responsable", a.get("responsable"))
        _docx_para(doc, "Plazo", a.get("plazo"))
        _docx_para(doc, "KPI", a.get("kpi"))
        _docx_para(doc, "Recursos", a.get("recursos"))
        _docx_para(doc, "Estado", a.get("estado"))
    if n == 0:
        doc.add_paragraph("Sin acciones registradas.")

    doc.add_heading("5. Evaluación e informes", level=1)
    _docx_para(doc, "Indicadores del plan de acción (KPI)", rend.get("kpis"))
    _docx_para(doc, "Frecuencia", rend.get("frecuencia_evaluacion"))
    _docx_para(doc, "Informes al comité", rend.get("informes_comite"))

    doc.add_heading("6. Recursos humanos y voluntarios", level=1)
    _docx_para(doc, "Roles clave", rrhh.get("roles_clave"))
    _docx_para(doc, "Brechas de formación", rrhh.get("brechas_formacion"))
    _docx_para(doc, "Reclutamiento", rrhh.get("reclutamiento"))
    _docx_para(doc, "Necesidades de voluntariado", vol.get("necesidades"))
    _docx_para(doc, "Motivaciones", vol.get("motivaciones"))
    _docx_para(doc, "Formación", vol.get("formacion"))
    _docx_para(doc, "Reconocimiento", vol.get("reconocimiento"))

    doc.add_heading("7. Primer proyecto (opcional)", level=1)
    _docx_para(doc, "Proyecto", pg.get("nombre"))
    _docx_para(doc, "Objetivo", pg.get("objetivo"))
    _docx_para(doc, "Contribuye al PEI en", pg.get("vinculo_estrategico"))
    _docx_para(doc, "Criterios de éxito", pg.get("criterios_exito"))

    doc.add_heading("8. Actividades de ejecución", level=1)
    n_act = 0
    for a in payload.get("actividades") or []:
        if not str(a.get("titulo") or "").strip():
            continue
        n_act += 1
        doc.add_heading(f"{n_act}. {a.get('titulo')}", level=2)
        _docx_para(doc, "Prioridad", a.get("prioridad"))
        _docx_para(doc, "Objetivo", a.get("objetivo"))
        _docx_para(doc, "Responsable", a.get("responsable"))
        _docx_para(doc, "Estado", a.get("estado"))
        _docx_para(doc, "KPI", a.get("kpi_nombre"))
        _docx_para(doc, "Meta", a.get("meta"))
        _docx_para(doc, "Avance", a.get("avance"))
    if n_act == 0:
        doc.add_paragraph("Sin actividades registradas.")

    surveys = list_configured_surveys(payload)
    if surveys:
        doc.add_heading("10. Encuestas de consulta", level=1)
        for s in surveys:
            _docx_para(doc, s["etiqueta"], s["url"])
            if s.get("destinatarios"):
                _docx_para(doc, "Destinatarios", s["destinatarios"])

    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Generado con Rumbo Deporte · Manual de Administración Deportiva, COI (2020), Unidades 53–57."
    )
    fr.italic = True
    fr.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
